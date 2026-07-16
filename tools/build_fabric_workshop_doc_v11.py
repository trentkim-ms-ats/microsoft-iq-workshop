from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

out_path = '/Users/hyungilkim/Documents/Data Platform Workshop/Fabric_Ontology_AI_Workshop_Detailed_Plan_v1.1.docx'

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.font.size = Pt(10.5)

title = doc.add_paragraph('Microsoft Fabric Ontology(Preview) 기반 AI 데이터 준비 워크숍 상세 개발 계획')
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
tr = title.runs[0]
tr.bold = True
tr.font.size = Pt(20)
tr.font.name = 'Arial'
tr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph('문서 버전: v1.1')
doc.add_paragraph('작성일: 2026-07-08')
doc.add_paragraph('대상: 워크숍 준비 인원 2명 (기술 리드 1, 콘텐츠/운영 리드 1)')
doc.add_paragraph('')

h = doc.add_paragraph('1. 워크숍 목적 및 성과 목표')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)
for g in [
    'Ontology(Preview)를 활용하여 AI 활용 가능한 데이터 구조(엔터티/관계/용어)를 설계하고 정합성을 확보한다.',
    'Fabric 내 데이터 품질 검증 기준(정확성, 완전성, 일관성, 적시성)을 실습으로 체득한다.',
    '준비된 데이터를 AI 검색/질의/생성 시나리오에 연결하여 실제 비즈니스 질문에 답할 수 있는 파이프라인을 완성한다.',
]:
    doc.add_paragraph(g, style='List Bullet')

p = doc.add_paragraph('성과 지표(KPI):')
p.runs[0].bold = True
for k in [
    '참가자 80% 이상이 Ontology 엔터티-관계 모델을 직접 작성 및 검증 완료',
    '참가자 80% 이상이 품질 검증 체크리스트 8개 항목 중 6개 이상 통과',
    '참가자 70% 이상이 AI 질의 시나리오 3개(검색/질의/생성) 모두 실행 성공',
]:
    doc.add_paragraph(k, style='List Bullet')

h = doc.add_paragraph('2. 1일(8시간) 상세 어젠다')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)

a = doc.add_table(rows=1, cols=5)
a.style = 'Table Grid'
hdr = a.rows[0].cells
hdr[0].text = '시간'
hdr[1].text = '세션'
hdr[2].text = '내용'
hdr[3].text = '방식'
hdr[4].text = '산출물'
agenda_rows = [
    ('09:00-09:20 (20분)', '오프닝/환경체크', '목표 공유, 계정/권한/데이터 연결 확인', '안내+점검', '접속 완료 체크'),
    ('09:20-10:00 (40분)', '개념 세션', 'Fabric + Ontology(Preview) 핵심 개념, 실습 흐름 설명', '강의', '실습 맵 이해'),
    ('10:00-11:30 (90분)', '트랙 1-1 데이터 준비', '원천 데이터 탐색, 스키마 정리, 엔터티 후보 도출', '실습', '엔터티 후보표'),
    ('11:30-11:40 (10분)', '휴식', 'Break', '휴식', '-'),
    ('11:40-12:30 (50분)', '트랙 1-2 Ontology 설계', '엔터티/관계/속성 정의, 매핑 규칙 작성', '실습', 'Ontology 초안'),
    ('12:30-13:30 (60분)', '점심', 'Lunch', '휴식', '-'),
    ('13:30-14:40 (70분)', '트랙 2-1 품질 규칙', '정확성/완전성/일관성/유효성/중복성/무결성 규칙 설계', '실습', '품질 규칙표'),
    ('14:40-15:20 (40분)', '트랙 2-2 검증 실행', '검증 쿼리 실행, 오류 분류, 품질 점수 산출', '실습', '품질 리포트'),
    ('15:20-15:30 (10분)', '휴식', 'Break', '휴식', '-'),
    ('15:30-16:30 (60분)', '트랙 3-1 AI 검색/질의', '준비 데이터 기반 검색/질의 시나리오 실행', '실습', '검색·질의 결과표'),
    ('16:30-17:10 (40분)', '트랙 3-2 AI 생성/평가', '요약/인사이트 생성, 환각/근거 연결성 평가', '실습', '생성 결과+평가표'),
    ('17:10-17:40 (30분)', '통합 미니 프로젝트', '3개 트랙 end-to-end 실행', '팀 실습', '최종 시나리오 결과'),
    ('17:40-18:00 (20분)', '리뷰/클로징', '발표, 피드백, 다음 액션 정리', '발표+Q&A', '개선 포인트'),
]
for row in agenda_rows:
    cells = a.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

p = doc.add_paragraph('트랙별 시간 합계:')
p.runs[0].bold = True
for s in [
    '트랙 1(데이터 준비 + Ontology): 140분',
    '트랙 2(검증/품질): 110분',
    '트랙 3(AI 활용): 100분',
    '통합 프로젝트: 30분',
    '오프닝/개념/클로징: 80분',
    '휴식/점심: 80분',
    '총합: 480분 (8시간)',
]:
    doc.add_paragraph(s, style='List Bullet')

h = doc.add_paragraph('3. 실습 트랙 상세 설계')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)

th = doc.add_paragraph('트랙 1. Fabric에서 Ontology(Preview) 기반 데이터 준비 (권장 140분)')
th.runs[0].bold = True
for line in [
    '목표: 비정형/정형 원천 데이터를 AI 친화적 구조로 정제하고 Ontology 엔터티/관계/속성을 설계한다.',
    '사전 준비물: 샘플 데이터(고객/주문/제품), Lakehouse, Notebook, 권한 템플릿',
    '산출물: Ontology 모델 초안, 엔터티 매핑표, 관계 정의서',
]:
    doc.add_paragraph(line, style='List Bullet')

t1 = doc.add_table(rows=1, cols=4)
t1.style = 'Table Grid'
hh = t1.rows[0].cells
hh[0].text = '단계'
hh[1].text = '시간'
hh[2].text = '준비 내용'
hh[3].text = '개발/실습 내용'
for r in [
    ('1) 요구사항 정리', '10분', '핵심 비즈니스 질문 도출', '도메인 질문 5개 정리'),
    ('2) 데이터 탐색/프로파일링', '30분', '결측/중복/이상 탐지 항목', 'Notebook 기반 프로파일링'),
    ('3) 표준 스키마 설계', '30분', '키/타입/코드 규칙', '표준화 로직 반영'),
    ('4) Ontology 엔터티/관계 설계', '40분', '엔터티·관계 정의 템플릿', '모델링 및 매핑'),
    ('5) 1차 검증', '30분', '무결성 체크리스트', '관계/참조 무결성 검증'),
]:
    c = t1.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

th = doc.add_paragraph('트랙 2. 준비 데이터 검증 및 품질 확인 (권장 110분)')
th.runs[0].bold = True
for line in [
    '목표: AI 입력 전에 품질 게이트를 통과하도록 검증 규칙과 점수화를 구성한다.',
    '사전 준비물: 품질 체크 규칙표, 검증 쿼리 템플릿, 샘플 오류 데이터셋',
    '산출물: 품질 검증 리포트, 오류 유형별 조치안, 품질 점수 대시보드',
]:
    doc.add_paragraph(line, style='List Bullet')
for c in [
    '정확성, 완전성, 일관성, 유효성, 중복성, 참조무결성, 적시성, 추적성',
    '품질 임계치 미달 항목에 대한 우선순위 조치안 수립',
]:
    doc.add_paragraph(c, style='List Number')

th = doc.add_paragraph('트랙 3. AI 활용(검색/질의/생성 시나리오) (권장 100분)')
th.runs[0].bold = True
for line in [
    '목표: 준비된 Ontology/검증 데이터로 AI 활용 시나리오를 실행하고 응답 품질을 평가한다.',
    '사전 준비물: 프롬프트 템플릿, 검색/질의 연결 설정, 평가 루브릭',
    '산출물: 시나리오별 프롬프트 세트, 결과 비교표, 개선 포인트',
]:
    doc.add_paragraph(line, style='List Bullet')

s = doc.add_table(rows=1, cols=4)
s.style = 'Table Grid'
sh = s.rows[0].cells
sh[0].text = '시나리오'
sh[1].text = '목표'
sh[2].text = '실습 내용'
sh[3].text = '평가 기준'
for r in [
    ('검색', '필요 데이터 탐색 정확도 향상', '도메인 용어 기반 검색 질의 구성', 'Top-k 적중률, 재현율'),
    ('질의', '비즈니스 질문에 구조화된 답변 생성', 'Ontology 관계 기반 질의 설계', '정답률, 근거 연결성'),
    ('생성', '요약/인사이트 생성', '검증 통과 데이터로 생성', '환각율, 실행가능성'),
]:
    c = s.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

h = doc.add_paragraph('4. 2인 준비체계 상세 (WBS)')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)

w = doc.add_table(rows=1, cols=5)
w.style = 'Table Grid'
wh = w.rows[0].cells
wh[0].text = 'Task'
wh[1].text = '준비/개발 내용'
wh[2].text = '담당'
wh[3].text = '공수(인시)'
wh[4].text = '완료 기준'
for r in [
    ('목표/범위 확정', '학습목표, 대상수준, KPI 정의', '공동', '16', '승인된 커리큘럼 1차안'),
    ('환경/권한 설계', '워크스페이스, 권한, 라이선스 점검', 'A', '24', '참가자 계정 점검표'),
    ('트랙1 개발', 'Ontology 실습 데이터/노트북/가이드', 'A', '40', '드라이런 통과'),
    ('트랙2 개발', '품질 규칙/검증쿼리/리포트 템플릿', 'A+B', '28', '품질 점수 산출 확인'),
    ('트랙3 개발', '검색/질의/생성 시나리오/평가표', 'A+B', '36', '시나리오 3종 성공'),
    ('교안/랩가이드', '슬라이드, 단계별 실습서, FAQ', 'B', '28', '수강생 리뷰 통과'),
    ('자동화/리셋', '환경 초기화, 데이터 리셋 스크립트', 'A', '16', '재실행 30분 내 복구'),
    ('드라이런 2회', '시간측정, 난이도 보정, 이슈수정', '공동', '24', '8시간 타임박스 준수'),
    ('운영체계/백업', '운영 R&R, 장애 대응 시나리오', 'B', '12', '운영 체크리스트 확정'),
    ('버퍼', 'Preview 기능 변경 대응', '공동', '20', '대체 실습 경로 확보'),
]:
    c = w.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

h = doc.add_paragraph('5. 리스크 및 대응')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)
for r in [
    'Preview 기능 변경: 대체 실습 스텝/사전 캡처본 확보',
    '권한/접속 오류: D-3 사전 점검 및 당일 헬프데스크 고정',
    '시간 초과: 필수(70%) + 선택(30%) 모듈화',
    'AI 편차: 평가 루브릭 + 정답 예시(근거 포함) 제공',
]:
    doc.add_paragraph(r, style='List Bullet')

h = doc.add_paragraph('6. 총 준비 공수 요약')
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)
doc.add_paragraph('총 준비 공수: 약 204 인시 (±20%)')
doc.add_paragraph('2명 기준 예상 캘린더: 약 4주 (버퍼 포함 4~5주)')

doc.save(out_path)
print(out_path)
