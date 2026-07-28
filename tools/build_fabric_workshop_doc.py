from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

out_path = '/Users/hyungilkim/Documents/Data Platform Workshop/Fabric_Ontology_AI_Workshop_Detailed_Plan.docx'

doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Arial'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.font.size = Pt(10.5)

# Title
title = doc.add_paragraph('Microsoft Fabric Ontology(Preview) 기반 AI 데이터 준비 워크숍 상세 개발 계획')
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
r = title.runs[0]
r.bold = True
r.font.size = Pt(20)
r.font.name = 'Arial'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph('문서 버전: v1.0')
doc.add_paragraph('작성일: 2026년 7월')
doc.add_paragraph('대상: 워크숍 준비 인원 2명 (기술 리드 1, 콘텐츠/운영 리드 1)')

doc.add_paragraph('')

h1 = doc.add_paragraph('1. 워크숍 목적 및 성과 목표')
h1.runs[0].bold = True
h1.runs[0].font.size = Pt(14)

goals = [
    'Ontology(Preview)를 활용하여 AI 활용 가능한 데이터 구조(엔터티/관계/용어)를 설계하고 정합성을 확보한다.',
    'Fabric 내 데이터 품질 검증 기준(정확성, 완전성, 일관성, 적시성)을 실습으로 체득한다.',
    '준비된 데이터를 AI 검색/질의/생성 시나리오에 연결하여 실제 비즈니스 질문에 답할 수 있는 파이프라인을 완성한다.',
]
for g in goals:
    p = doc.add_paragraph(g, style='List Bullet')

p = doc.add_paragraph('성과 지표(KPI):')
p.runs[0].bold = True
for k in [
    '참가자 80% 이상이 Ontology 엔터티-관계 모델을 직접 작성 및 검증 완료',
    '참가자 80% 이상이 품질 검증 체크리스트 8개 항목 중 6개 이상 통과',
    '참가자 70% 이상이 AI 질의 시나리오 3개(검색/질의/생성) 모두 실행 성공',
]:
    doc.add_paragraph(k, style='List Bullet')

h2 = doc.add_paragraph('2. 워크숍 운영 가정')
h2.runs[0].bold = True
h2.runs[0].font.size = Pt(14)
for a in [
    '워크숍 형태: 1일 6시간(이론 1.5시간 + 실습 4시간 + Q&A 0.5시간)',
    '참가자 규모: 15~25명',
    '대상 수준: 데이터/SQL 기초 보유, Fabric 기초 경험 권장',
    '사전 준비: Fabric 접근 권한, 샘플 데이터셋 배포, 실습 계정 점검',
]:
    doc.add_paragraph(a, style='List Bullet')

h3 = doc.add_paragraph('3. 실습 트랙 상세 설계')
h3.runs[0].bold = True
h3.runs[0].font.size = Pt(14)

# Track 1
th = doc.add_paragraph('트랙 1. Fabric에서 Ontology(Preview) 기반 데이터 준비 (권장 100분)')
th.runs[0].bold = True

for line in [
    '목표: 비정형/정형 원천 데이터를 AI 친화적 구조로 정제하고 Ontology 엔터티/관계/속성을 설계한다.',
    '사전 준비물: 샘플 데이터(고객/주문/제품), Lakehouse, Notebook, 기본 권한 템플릿',
    '산출물: Ontology 모델 초안, 엔터티 매핑표, 관계 정의서',
]:
    doc.add_paragraph(line, style='List Bullet')

p = doc.add_paragraph('세부 실습 단계(예시):')
p.runs[0].bold = True
steps1 = [
    ('1) 데이터 도메인 브리핑 및 요구사항 정리', '10분', '고객-주문-제품 핵심 질문 5개 도출'),
    ('2) 원천 데이터 탐색/프로파일링', '20분', '결측치/중복/코드값 이상 항목 식별'),
    ('3) 표준 스키마 설계', '20분', '공통 키, 타입, 표준 코드 규칙 확정'),
    ('4) Ontology 엔터티/관계 정의', '30분', '엔터티 6~10개, 관계 8~15개 설계'),
    ('5) 매핑 및 1차 검증', '20분', '관계 무결성/참조 무결성 체크'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '단계'
hdr[1].text = '시간'
hdr[2].text = '준비 내용'
hdr[3].text = '개발/실습 내용'
for s, t, d in steps1:
    row = table.add_row().cells
    row[0].text = s
    row[1].text = t
    row[2].text = d
    row[3].text = 'Notebook/SQL 기반 정제 및 Ontology 매핑 실습'

# Track 2
th2 = doc.add_paragraph('\n트랙 2. 준비 데이터 검증 및 품질 확인 (권장 80분)')
th2.runs[0].bold = True
for line in [
    '목표: AI 입력 전에 품질 게이트를 통과하도록 데이터 품질 기준과 자동 점검 항목을 만든다.',
    '사전 준비물: 품질 체크 규칙표, 검증 쿼리 템플릿, 샘플 오류 데이터셋',
    '산출물: 품질 검증 리포트, 오류 유형별 조치안, 품질 점수 대시보드',
]:
    doc.add_paragraph(line, style='List Bullet')

p = doc.add_paragraph('핵심 품질 점검 항목(최소 8개):')
p.runs[0].bold = True
checks = [
    '정확성: 기준 마스터와 값 일치율',
    '완전성: 필수 컬럼 결측률',
    '일관성: 동일 엔터티 속성 충돌률',
    '유효성: 형식/범위 규칙 위반률',
    '중복성: 키 중복 건수',
    '참조무결성: 부모-자식 키 미스매치',
    '적시성: 최신성 SLA 충족률',
    '추적성: 원천-가공 lineage 확인 가능 여부',
]
for c in checks:
    doc.add_paragraph(c, style='List Number')

# Track 3
th3 = doc.add_paragraph('\n트랙 3. AI 활용(검색/질의/생성 시나리오) (권장 90분)')
th3.runs[0].bold = True
for line in [
    '목표: 준비된 Ontology/검증 데이터로 AI 활용 시나리오를 구현하고 응답 품질을 평가한다.',
    '사전 준비물: 프롬프트 템플릿, 검색 인덱스 또는 질의 연결 설정, 평가 루브릭',
    '산출물: 시나리오별 프롬프트 세트, 응답 결과 비교표, 개선 포인트 목록',
]:
    doc.add_paragraph(line, style='List Bullet')

scen_table = doc.add_table(rows=1, cols=4)
scen_table.style = 'Table Grid'
h = scen_table.rows[0].cells
h[0].text = '시나리오'
h[1].text = '목표'
h[2].text = '실습 내용'
h[3].text = '평가 기준'
rows = [
    ('검색', '필요 데이터 탐색 정확도 향상', '도메인 용어 기반 검색 질의 구성', 'Top-k 적중률, 재현율'),
    ('질의', '비즈니스 질문에 구조화된 답변 생성', 'Ontology 관계를 활용한 질의 설계', '정답률, 근거 연결성'),
    ('생성', '요약/인사이트 문장 생성', '품질 검증 통과 데이터만 사용해 생성', '환각율, 실행가능성'),
]
for r0, r1, r2, r3 in rows:
    rw = scen_table.add_row().cells
    rw[0].text = r0
    rw[1].text = r1
    rw[2].text = r2
    rw[3].text = r3

h4 = doc.add_paragraph('\n4. 2인 준비체계 상세 (WBS 기반)')
h4.runs[0].bold = True
h4.runs[0].font.size = Pt(14)

wbs = doc.add_table(rows=1, cols=5)
wbs.style = 'Table Grid'
wh = wbs.rows[0].cells
wh[0].text = 'Task'
wh[1].text = '준비/개발 내용'
wh[2].text = '담당'
wh[3].text = '공수(인시)'
wh[4].text = '완료 기준'

wbs_rows = [
    ('목표/범위 확정', '학습목표, 대상수준, KPI 정의', '공동', '16', '승인된 커리큘럼 1차안'),
    ('환경/권한 설계', '워크스페이스, 권한, 라이선스 점검', 'A', '24', '참가자 계정 점검표'),
    ('트랙1 개발', 'Ontology 실습 데이터/노트북/가이드', 'A', '40', '드라이런 통과'),
    ('트랙2 개발', '품질 규칙/검증쿼리/리포트 템플릿', 'A+B', '28', '품질 점수 산출 확인'),
    ('트랙3 개발', '검색/질의/생성 시나리오 및 평가표', 'A+B', '36', '시나리오 3종 성공'),
    ('교안/랩가이드', '슬라이드, 단계별 실습서, FAQ', 'B', '28', '수강생 리뷰 통과'),
    ('자동화/리셋', '환경 초기화, 데이터 리셋 스크립트', 'A', '16', '재실행 30분 내 복구'),
    ('드라이런 2회', '시간측정, 난이도 보정, 이슈수정', '공동', '24', '총시간 6시간 이내'),
    ('운영체계/백업', '운영 R&R, 장애 대응 시나리오', 'B', '12', '운영 체크리스트 확정'),
    ('버퍼', 'Preview 이슈 대응', '공동', '20', '대체 실습 경로 확보'),
]
for a, b, c, d, e in wbs_rows:
    rr = wbs.add_row().cells
    rr[0].text = a
    rr[1].text = b
    rr[2].text = c
    rr[3].text = d
    rr[4].text = e

h5 = doc.add_paragraph('\n5. 일정 계획 (2명, 4주)')
h5.runs[0].bold = True
h5.runs[0].font.size = Pt(14)
for wk in [
    '1주차: 범위 확정, 환경 구축, 트랙1 설계 시작 (48~56 인시)',
    '2주차: 트랙1 완성, 트랙2/3 1차 개발 (52~60 인시)',
    '3주차: 교안/핸드북 완성, 자동화, 드라이런 1차 (48~56 인시)',
    '4주차: 드라이런 2차, 최종 수정, 운영 리허설 (40~52 인시)',
]:
    doc.add_paragraph(wk, style='List Bullet')

h6 = doc.add_paragraph('\n6. 리스크 및 대응')
h6.runs[0].bold = True
h6.runs[0].font.size = Pt(14)
for r in [
    'Preview 기능 변경 리스크: 핵심 화면/기능 캡처본 + 대체 실습 스텝 준비',
    '권한/접속 오류 리스크: 사전 점검일(D-3) 운영, 당일 헬프데스크 담당 고정',
    '시간 초과 리스크: 필수 실습(70%) + 선택 실습(30%)으로 모듈화',
    'AI 결과 편차 리스크: 평가 루브릭과 정답 예시(근거 포함) 제공',
]:
    doc.add_paragraph(r, style='List Bullet')

h7 = doc.add_paragraph('\n7. 당일 운영 체크리스트')
h7.runs[0].bold = True
h7.runs[0].font.size = Pt(14)
for c in [
    'D-1: 계정/권한/데이터 리셋 최종 점검 완료',
    'T-30분: 실습 환경 로그인 및 샘플 쿼리 실행 확인',
    '진행 중: 30분 단위 진도/이탈자 수집, 보조강사 순회',
    '종료 직전: KPI 체크 설문, 실습 산출물 제출 확인',
    '종료 후: 이슈 로그 정리, 차수 개선안 업데이트',
]:
    doc.add_paragraph(c, style='List Bullet')

h8 = doc.add_paragraph('\n8. 총 준비 공수 요약')
h8.runs[0].bold = True
h8.runs[0].font.size = Pt(14)
doc.add_paragraph('총 준비 공수: 약 204 인시 (±20%)')
doc.add_paragraph('2명 기준 예상 캘린더: 약 4주 (버퍼 포함 4~5주)')

doc.save(out_path)
print(out_path)
